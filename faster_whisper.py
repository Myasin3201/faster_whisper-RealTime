import sounddevice as sd
import queue
import threading
import numpy as np
from faster_whisper import WhisperModel
from docx import Document
import os
import keyboard

samplerate = 16000
block_duration = 0.5
chunk_duration = 2
channels = 1
running = True

output_file = "transcript.docx"
if not os.path.exists(output_file):
    Document().save(output_file)

frames_per_block = int(samplerate * block_duration)
frames_per_chunk = int(samplerate * chunk_duration)

audio_queue = queue.Queue()
audio_buffer = []

model = WhisperModel("small", compute_type="int8")

def stop_listener():
    global running
    keyboard.wait("esc")
    running = False

def audio_callback(in_data, frames, time, status):
    if status:
        print(status)
    audio_queue.put(in_data.copy())

def recorder():
    with sd.InputStream(samplerate=samplerate , channels=channels ,
                        callback=audio_callback , blocksize=frames_per_block):
        print("listening ...")
        while running:
            sd.sleep(100)
doc = Document(output_file)
def transcriber():
    global audio_buffer  , doc
    while running:
        block = audio_queue.get()
        audio_buffer.append(block)

        total_frames = sum(len(b) for b in audio_buffer)
        if total_frames >= frames_per_chunk:

            audio_data = np.concatenate(audio_buffer)[:frames_per_chunk]
            # audio_buffer = []
            audio_buffer = audio_buffer[frames_per_chunk:]

            # full_audio = np.concatenate(audio_buffer)
            # audio_data = full_audio[:frames_per_chunk]
            # leftover = full_audio[frames_per_chunk:]
            # audio_buffer = [leftover] if len(leftover) > 0 else []

            audio_data = audio_data.flatten().astype(np.float32)

            volume = np.abs(audio_data).mean()
            if volume < 0.01:
                continue

            segments , _ = model.transcribe(
                audio_data,
                language="en",
                beam_size=1  ,
                vad_filter=True,
                vad_parameters=dict(min_silence_duration_ms = 500)
            )
            has_text = False
            for segment in segments:
                print(f'{segment.text}')
                doc.add_paragraph(segment.text)
                has_text = True
            if has_text:
                doc.save(output_file)

    doc.save(output_file)

threading.Thread(target=stop_listener, daemon=True).start()
threading.Thread(target = recorder , daemon = True).start()
transcriber()